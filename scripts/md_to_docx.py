from docx import Document
from docx.shared import Pt, RGBColor, Inches
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

SRC = r'c:\Users\rosej\Downloads\Bloom\docs\BLOOM_SYSTEM_OVERVIEW_DEMO_QA.docx'
OUT = SRC

with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.readlines()

def clean(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()

in_code = False
code_lines = []
table_rows = []

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Code block toggle
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            p.paragraph_format.left_indent = Inches(0.4)
            code_lines = []
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Table rows
    if line.startswith('|'):
        cells = [c.strip() for c in line.strip('|').split('|')]
        is_separator = all(re.match(r'^[-: ]+$', c) for c in cells if c)
        if not is_separator:
            table_rows.append(cells)
        i += 1
        next_is_table = i < len(lines) and lines[i].startswith('|')
        if not next_is_table and table_rows:
            max_cols = max(len(r) for r in table_rows)
            t = doc.add_table(rows=len(table_rows), cols=max_cols)
            t.style = 'Table Grid'
            for r_idx, row in enumerate(table_rows):
                for c_idx, cell in enumerate(row):
                    if c_idx < max_cols:
                        tc = t.rows[r_idx].cells[c_idx]
                        tc.text = clean(cell)
                        if r_idx == 0:
                            for para in tc.paragraphs:
                                for run in para.runs:
                                    run.bold = True
            table_rows = []
            doc.add_paragraph()
        continue

    # Headings
    if line.startswith('#### '):
        doc.add_heading(clean(line[5:]), level=4)
    elif line.startswith('### '):
        doc.add_heading(clean(line[4:]), level=3)
    elif line.startswith('## '):
        doc.add_heading(clean(line[3:]), level=2)
    elif line.startswith('# '):
        doc.add_heading(clean(line[2:]), level=1)
    elif line.strip() == '---':
        doc.add_paragraph('_' * 60)
    elif line.startswith('- Q:'):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(clean(line[2:]))
        run.bold = True
    elif line.strip().startswith('- A:'):
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(clean(line.strip()[4:]))
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(clean(line[2:]))
    elif re.match(r'^\d+\. ', line):
        p = doc.add_paragraph(style='List Number')
        p.add_run(clean(re.sub(r'^\d+\. ', '', line)))
    elif line.startswith('> '):
        p = doc.add_paragraph(clean(line[2:]))
        p.paragraph_format.left_indent = Inches(0.4)
        for run in p.runs:
            run.italic = True
    elif line.strip() == '':
        doc.add_paragraph()
    else:
        if line.strip():
            doc.add_paragraph(clean(line))

    i += 1

doc.save(OUT)
print('Saved:', OUT)
